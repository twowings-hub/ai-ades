# -*- coding: utf-8 -*-
"""
AI-ADES Phase 5·6 작업 정리 + 고객사 미팅 협의/확인사항 문서 생성기
- 동일 내용을 Word(.docx)와 HTML 두 형식으로 출력한다 (내용 동기화 보장)
출력: docs/Phase5-6_고객사미팅_정리.docx / .html
"""
import os
import html as _html

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCS_DIR = os.path.join(os.path.dirname(__file__), "..", "docs")
BASENAME = "Phase5-6_고객사미팅_정리"
TODAY = "2026-06-13"

ACCENT = "2563EB"

# ------------------------------------------------------------
# 문서 내용 정의 (블록 리스트) — 한 곳에서 정의해 두 형식이 공유
#   ("h1", text) / ("h2", text) / ("p", text)
#   ("ul", [items]) / ("check", [items]) / ("table", [[...rows...]] with header row)
#   ("note", text) — 강조 박스
# ------------------------------------------------------------
BLOCKS = [
    ("title", "AI-ADES Phase 5·6 작업 정리 및 고객사 미팅 협의·확인사항"),
    ("subtitle", "CO₂ 레이저 Glass+Film 적층소재 절단 — 최적 레이저조건 자동탐색 시스템"),
    ("meta", f"작성일: {TODAY}  ·  대상: 고객사 설치/연동 협의  ·  문서버전: v1.0"),

    ("h1", "1. 문서 목적"),
    ("p", "본 문서는 (1) 앞으로 남은 Phase 5(고객사 설치 테스트)·Phase 6(설비 연동)의 "
          "할 일을 정리하고, (2) 고객사 미팅에서 반드시 협의·확인해야 할 사항을 체크리스트로 "
          "제공하기 위한 것이다. 미팅 시 본 체크리스트를 함께 보며 항목별로 확정한다."),

    ("h1", "2. 현재 진행 상황 요약"),
    ("table", [
        ["Phase", "내용", "상태"],
        ["Phase 0", "환경 세팅, Docker, DB 스키마", "완료"],
        ["Phase 1", "Data Prep Agent, Excel 파싱", "완료"],
        ["Phase 2", "Modeling Agent, XGBoost + SHAP", "완료"],
        ["Phase 3", "Execution Agent, Bayesian Opt, 승인/레시피, Admin", "완료"],
        ["Phase 4", "React UI + Admin Console (UI 마무리 중)", "진행 중"],
        ["Phase 5", "고객사 워크스테이션 2대 설치 테스트", "대기 (본 문서)"],
        ["Phase 6", "CO₂ 레이저 설비 연동 + 실데이터 학습", "대기 (본 문서)"],
    ]),
    ("note", "Phase 5·6의 상당 부분은 고객사 하드웨어/현장이 있어야 검증 가능하다. "
             "개발 PC(Surface)에서 미리 만들 수 있는 것은 배포 패키징(아래 3.1)에 한정된다."),

    ("h1", "3. Phase 5 — 고객사 설치 테스트"),
    ("p", "목표: 고객 워크스테이션 2대 설치 + GPU 환경 검증. "
          "완료 기준은 고객사 하드웨어에서 E2E(실험조건→AI제안→승인→결과보고) 루프 3회 완주."),

    ("h2", "3.1 개발측 사전 준비물 (지금 개발 PC에서 제작 가능)"),
    ("check", [
        "docker-compose.prod.yml — 운영용 compose (GPU_ENABLED 옵션 포함)",
        "install.bat — Windows 원클릭 설치 스크립트",
        "env.check.py — CUDA 인식 / 포트 충돌 / 메모리 사전 점검 스크립트",
        "오프라인 설치 대비: Docker 이미지 사전 빌드·반입 패키지(인터넷 차단 환경 시)",
        "설치 가이드 문서(설치 순서·트러블슈팅)",
    ]),

    ("h2", "3.2 현장 설치 작업 (고객사 하드웨어 필요)"),
    ("p", "[설계해석 워크스테이션: Xeon W5-2545 + RTX PRO 4000 Blackwell]"),
    ("check", [
        "NVIDIA Driver 570+ 설치 확인",
        "Docker GPU 패스스루 (NVIDIA Container Toolkit) 구성",
        "PyTorch GPU 모드 재구성 (LSTM 학습 속도 향상)",
        "docker-compose.prod.yml 전체 기동",
        "GPU 인식 확인 (nvidia-smi / 컨테이너 내부)",
    ]),
    ("p", "[CPU 특화 워크스테이션: ULTRA 9 285]"),
    ("check", [
        "Kafka / InfluxDB / PostgreSQL 전담 구동",
        "두 워크스테이션 간 내부 네트워크 연결 확인",
        "Excel POC 데이터로 Phase 1~4 재검증",
    ]),

    ("h2", "3.3 완료 기준"),
    ("ul", ["고객사 H/W에서 E2E 루프 3회 완주 성공"]),

    ("h1", "4. Phase 6 — 설비 연동 (고객사 현장)"),
    ("p", "목표: 실제 CO₂ 레이저 설비 연동 + Plasma 시계열 학습 + 실시간 운영. "
          "완료 기준은 실설비에서 Auto DOE 5~7회 내 OK 수렴 실증 및 납품 시연 완료."),

    ("h2", "4.1 현장 실사 확인 사항"),
    ("check", [
        "레이저 설비 통신 인터페이스 종류 (OPC-UA / RS-232 / Ethernet / 설비 전용 API)",
        "실시간으로 읽을 수 있는 파라미터 목록 (Speed/Defocus/Frequency/Power 실시간 노출 여부)",
        "가공 완료 후 Kerf/Depth 자동 측정 장치 유무 (없으면 운영자 수동 입력 UI 사용)",
        "Plasma sensor ZIP 데이터 포맷 확인 (250kHz CSV: Index;Time;Area;Plasma;P-Raw;Temp;T-Raw;Refl;R-Raw)",
        "공장 네트워크 IP 대역 / 방화벽 포트 허용 정책",
    ]),

    ("h2", "4.2 개발 작업"),
    ("check", [
        "설비 통신 드라이버 개발 (OPC-UA 또는 해당 프로토콜)",
        "Plasma sensor ZIP 파서 → plasma_timeseries 테이블 적재",
        "LSTM-Autoencoder 학습 (Plasma 시계열 이상감지, M1/M2/Blank 구간 자동 분리, F1 ≥ 0.80 목표)",
        "실데이터 기반 모델 재학습 (POC Excel → 실설비 데이터 점진 교체)",
        "Airflow DAG: 50건 누적 시 자동 재학습 트리거",
        "24시간 연속 수집 테스트",
        "납품 시연 리허설 3회 → 최종 시연",
    ]),

    ("h2", "4.3 완료 기준"),
    ("ul", ["실설비 Auto DOE 5~7회 내 OK 수렴 실증", "납품 시연 완료"]),

    ("h1", "5. 고객사 현장 현황 파악(As-Is) 및 협의 체크리스트"),
    ("p", "미팅에서 아래 항목을 순서대로 확인하고, 각 항목의 응답·결정값을 비고란(현장 메모)에 기록한다. "
          "특히 5.1~5.2(설비)와 5.5(데이터 취득), 5.7(개발기간/운영), 5.8(발전계획)은 "
          "현장 실물을 보며 구체적으로 확인한다."),

    ("h2", "5.1 CO₂ 레이저 가공기 (설비 현황)"),
    ("check", [
        "제조사 / 모델명 / 연식 / 설치 대수",
        "레이저 종류·출력 사양 (CO₂ 출력 W, 발진 방식, 펄스/CW)",
        "제어기(컨트롤러) 종류·제조사 (PLC / CNC / 설비 전용 제어기)",
        "운전 파라미터 설정 방식 (HMI 수동 입력 / 레시피 / 외부 통신)",
        "현재 조건 탐색 방법 (수동 시행착오 횟수, 1회 가공·측정 소요 시간)",
        "설비 가동률 / 일일 가공 건수 / 24시간 가동 여부",
    ]),

    ("h2", "5.2 설비 인터페이스 / 연동 방식 (Phase 6 핵심)"),
    ("check", [
        "통신 인터페이스 종류 (OPC-UA / RS-232 / Ethernet/IP / Modbus / 설비 전용 API)",
        "제어 통신 규격 문서·SDK·태그맵 제공 가능 여부",
        "실시간 읽기 가능 파라미터 목록 및 갱신 주기 (Speed/Defocus/Frequency/Power)",
        "AI 제안 조건을 설비에 '직접 쓰기' 가능한가, 아니면 운영자 수동 입력인가 (제어 권한 범위)",
        "Kerf/Depth 자동 측정 장치 유무·종류 — 없으면 운영자 수동 입력(현 승인 화면 활용)",
        "Plasma sensor 장착 여부 / 데이터 실샘플 1건 확보 (포맷·컬럼 검증용)",
        "설비 비상정지·인터락 연동 필요 여부 및 안전 요구사항",
    ]),

    ("h2", "5.3 준비된 AI 실행 장비 (하드웨어 실물 확인)"),
    ("check", [
        "워크스테이션 종류·대수·실물 스펙 (CPU / GPU 모델 / RAM / 디스크)",
        "GPU 드라이버 버전(NVIDIA 570+), CUDA, OS(Windows 11) 버전",
        "관리자 권한 / Docker Desktop·NVIDIA Container Toolkit 설치 가능 여부",
        "설치 장소 환경 (설비와의 거리, 분진/온습도, 전원/UPS)",
        "설비-AI장비 물리적 연결 경로 (유선 LAN / 산업용 스위치 등)",
    ]),

    ("h2", "5.4 네트워크 / 보안"),
    ("check", [
        "사용 포트 충돌 점검: 8010, 8011, 8012, 5173, 8086, 5000, 3010, 9092, 5432, 6333, 11434",
        "두 워크스테이션 간 통신 가능 여부 (IP 대역/서브넷/방화벽)",
        "인터넷 연결 여부 (오프라인이면 Docker 이미지 사전 반입 필요)",
        "LLM 사용 정책: 로컬(Ollama)만 vs 외부 API 허용 — 가공데이터 외부 반출 가능 여부",
        "공장 보안 정책 / 계정·권한 / 데이터 보관·백업 위치 규정",
    ]),

    ("h2", "5.5 데이터 현황 및 학습 데이터 취득 방법"),
    ("check", [
        "기존 보유 실험 데이터 양·형식 (Excel / DB / 종이대장) 및 제공 가능 여부",
        "POC 외 추가 과거 데이터 규모 (조건-결과 쌍 건수)",
        "신규 데이터 취득 방식: 가공 1건당 자동 적재 가능? 운영자 수동 입력?",
        "결과 라벨링(판정) 주체·기준 (검사자 육안 / 측정장비 자동)",
        "측정 자동화 수준 (Kerf/Depth/Defect 측정의 자동·반자동·수동)",
        "데이터 품질·이상치 관리 기준, 측정 누락·재현성 이슈",
        "데이터 누적 속도 — 자동 재학습 트리거(50건) 도달 예상 주기",
    ]),

    ("h2", "5.6 소재 / 판정 기준 / 탐색 공간"),
    ("check", [
        "실소재 조합·두께 범위 (M1=Glass, M2=Film) — 학습 데이터 범위와 일치 여부",
        "판정 기준 확정: OK = 0μm 초과 ~ 25μm 이하 유지 여부, 미가공/과가공/NG·Defect 정의 합의",
        "탐색 공간 확정: speed[200,500,1000] / defocus[0~4] / frequency[100,200] / power[2.8~59.8] — 실설비와 일치 여부",
        "소재 종류 추가·변경 가능성 (향후 다른 소재 적용 계획)",
    ]),

    ("h2", "5.7 개발 기간 / 예산 / 운영 방법"),
    ("check", [
        "희망 개발 기간·납기 및 주요 마일스톤",
        "계약 범위 (설치 / 설비연동 / 유지보수 포함 여부) 및 예산 규모",
        "운영 주체: 고객 자체 운영 vs 위탁 운영",
        "운영 인력 수·교대, 가동 시간(24시간 여부), 사용 권한 레벨(운영자/관리자)",
        "Human-in-the-Loop 승인 주체 및 책임 범위",
        "유지보수·장애대응 SLA, 원격 지원 허용 여부",
        "운영자 교육·매뉴얼 제공 범위",
    ]),

    ("h2", "5.8 향후 시스템 발전 계획 (To-Be 로드맵)"),
    ("check", [
        "적용 확대 대상: 다른 소재/두께/제품군 추가",
        "다중 설비 동시 운영 / 생산라인 확장 계획",
        "모델 고도화: 실데이터 누적 재학습, 신규 센서(Plasma 등) 활용 범위",
        "상위 시스템 연동 (MES / ERP / 품질관리시스템)",
        "원격 모니터링·클라우드 확장, 다지점 운영 가능성",
        "장기 목표 (완전 무인 자동 최적화 수준, 적용 KPI)",
    ]),

    ("h2", "5.9 일정 / 납품 / 검수"),
    ("check", [
        "설치 일정 및 현장 방문 가능일 확정",
        "Phase 5 완료 기준(고객 H/W에서 E2E 3회 완주) 합의",
        "Phase 6 완료 기준(실설비 Auto DOE 5~7회 OK 수렴 + 시연) 합의",
        "납품 시연 일정·시나리오·참석자 확정",
        "검수 기준 및 인수인계 범위",
    ]),

    ("h1", "6. 미팅에서 반드시 결정해야 할 핵심 사항 (요약)"),
    ("table", [
        ["#", "결정 사항", "왜 중요한가"],
        ["1", "설비 통신 방식 (OPC-UA/RS-232/Ethernet/API) 및 문서·SDK 제공", "Phase 6 드라이버 개발 범위·기간을 좌우"],
        ["2", "AI 조건 설비 직접 쓰기 vs 운영자 수동 입력", "자동화 수준·안전 인터락 설계 결정"],
        ["3", "Kerf/Depth 자동 측정 유무", "결과 수집 자동화 가능 여부, UI 변경 여부"],
        ["4", "학습 데이터 취득 방식 (과거 보유량 + 신규 수집 자동화)", "모델 정확도·재학습 주기·자동화 범위 결정"],
        ["5", "AI 실행 장비 실물 스펙·GPU 환경", "설치 가능성·성능(LSTM 학습 속도) 좌우"],
        ["6", "LLM 외부 API 허용 여부", "데이터 보안 정책 / 로컬 Ollama 의존 여부"],
        ["7", "판정 기준·탐색 공간 실설비 일치", "모델 재학습·재보정 필요 범위"],
        ["8", "개발 기간 / 예산 / 운영 주체", "계약·납기·유지보수 스코프 확정"],
        ["9", "향후 발전 계획 우선순위", "확장성 고려한 아키텍처 설계 방향"],
        ["10", "오프라인 설치 여부 / 설치·시연 일정", "사전 패키지 준비 및 스케줄 확정"],
    ]),
    ("note", "위 항목은 미팅에서 결론이 나야 Phase 6 개발 착수가 가능하다. "
             "특히 1~4번(설비 연동 방식과 데이터 취득 체계)은 개발 공수·일정에 가장 큰 영향을 준다. "
             "8~9번(기간·예산·발전계획)은 아키텍처와 계약 범위를 좌우하므로 함께 확정한다."),
]


# ------------------------------------------------------------
# Word(.docx) 생성
# ------------------------------------------------------------
def build_docx(path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10.5)

    for kind, payload in BLOCKS:
        if kind == "title":
            p = doc.add_paragraph()
            run = p.add_run(payload)
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor.from_string(ACCENT)
        elif kind == "subtitle":
            p = doc.add_paragraph()
            run = p.add_run(payload)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor.from_string("555555")
        elif kind == "meta":
            p = doc.add_paragraph()
            run = p.add_run(payload)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string("888888")
            doc.add_paragraph()
        elif kind == "h1":
            doc.add_heading(payload, level=1)
        elif kind == "h2":
            doc.add_heading(payload, level=2)
        elif kind == "p":
            doc.add_paragraph(payload)
        elif kind == "ul":
            for item in payload:
                doc.add_paragraph(item, style="List Bullet")
        elif kind == "check":
            for item in payload:
                doc.add_paragraph(f"☐  {item}")
        elif kind == "note":
            p = doc.add_paragraph()
            run = p.add_run("※ " + payload)
            run.italic = True
            run.font.color.rgb = RGBColor.from_string("8A4B00")
        elif kind == "table":
            rows = payload
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Light Grid Accent 1"
            for r, row in enumerate(rows):
                for c, val in enumerate(row):
                    cell = table.cell(r, c)
                    cell.text = str(val)
                    if r == 0:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
            doc.add_paragraph()

    doc.save(path)


# ------------------------------------------------------------
# HTML 생성
# ------------------------------------------------------------
def build_html(path):
    def esc(s):
        return _html.escape(str(s))

    parts = [
        "<!DOCTYPE html>",
        "<html lang='ko'><head><meta charset='utf-8'>",
        "<meta name='viewport' content='width=device-width, initial-scale=1'>",
        "<title>AI-ADES Phase 5·6 고객사 미팅 정리</title>",
        "<style>",
        "body{font-family:'Malgun Gothic','맑은 고딕',sans-serif;color:#1f2329;",
        "max-width:900px;margin:0 auto;padding:32px 24px;line-height:1.7;}",
        "h1{font-size:20px;border-bottom:2px solid #2563eb;padding-bottom:6px;margin-top:36px;}",
        "h2{font-size:16px;color:#1d4ed8;margin-top:24px;}",
        ".title{font-size:26px;font-weight:700;color:#2563eb;margin-bottom:4px;}",
        ".subtitle{font-size:14px;color:#555;}",
        ".meta{font-size:12px;color:#888;margin-bottom:8px;}",
        "table{border-collapse:collapse;width:100%;margin:12px 0;font-size:14px;}",
        "th,td{border:1px solid #d0d3d8;padding:7px 10px;text-align:left;vertical-align:top;}",
        "th{background:#eef2ff;}",
        "ul{margin:8px 0;padding-left:22px;}",
        ".check{list-style:none;padding-left:4px;margin:8px 0;}",
        ".check li{margin:5px 0;}",
        ".check li:before{content:'\\2610';margin-right:8px;color:#2563eb;font-size:16px;}",
        ".note{background:#fff7ed;border-left:4px solid #f59e0b;padding:10px 14px;",
        "margin:14px 0;color:#8a4b00;font-size:13px;border-radius:0 6px 6px 0;}",
        "@media print{body{padding:0;}h1{page-break-after:avoid;}}",
        "</style></head><body>",
    ]

    for kind, payload in BLOCKS:
        if kind == "title":
            parts.append(f"<div class='title'>{esc(payload)}</div>")
        elif kind == "subtitle":
            parts.append(f"<div class='subtitle'>{esc(payload)}</div>")
        elif kind == "meta":
            parts.append(f"<div class='meta'>{esc(payload)}</div>")
        elif kind == "h1":
            parts.append(f"<h1>{esc(payload)}</h1>")
        elif kind == "h2":
            parts.append(f"<h2>{esc(payload)}</h2>")
        elif kind == "p":
            parts.append(f"<p>{esc(payload)}</p>")
        elif kind == "ul":
            parts.append("<ul>" + "".join(f"<li>{esc(i)}</li>" for i in payload) + "</ul>")
        elif kind == "check":
            parts.append("<ul class='check'>" + "".join(f"<li>{esc(i)}</li>" for i in payload) + "</ul>")
        elif kind == "note":
            parts.append(f"<div class='note'>※ {esc(payload)}</div>")
        elif kind == "table":
            rows = payload
            cells = []
            cells.append("<tr>" + "".join(f"<th>{esc(c)}</th>" for c in rows[0]) + "</tr>")
            for row in rows[1:]:
                cells.append("<tr>" + "".join(f"<td>{esc(c)}</td>" for c in row) + "</tr>")
            parts.append("<table>" + "".join(cells) + "</table>")

    parts.append("</body></html>")
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(parts))


if __name__ == "__main__":
    os.makedirs(DOCS_DIR, exist_ok=True)
    docx_path = os.path.abspath(os.path.join(DOCS_DIR, BASENAME + ".docx"))
    html_path = os.path.abspath(os.path.join(DOCS_DIR, BASENAME + ".html"))
    build_docx(docx_path)
    build_html(html_path)
    print("생성 완료:")
    print(" -", docx_path)
    print(" -", html_path)
